from io import BytesIO

import pandas as pd
from docx import Document
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def generate_pdf_report(report: dict):
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    elements = []

    summary = report.get("executive_summary", {})
    findings = report.get("findings", [])
    recommendations = report.get("recommendations", [])

    elements.append(
        Paragraph(
            "AI Compliance & Document Review Report",
            styles["Title"],
        )
    )
    elements.append(Spacer(1, 20))

    elements.append(
        Paragraph(
            f"<b>File:</b> {report.get('file_name', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(
        Paragraph(
            f"<b>Framework:</b> {report.get('framework', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(
        Paragraph(
            f"<b>Document Type:</b> {report.get('document_type', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(
        Paragraph(
            f"<b>Generated:</b> {report.get('created_at', 'N/A')}",
            styles["Normal"],
        )
    )
    elements.append(Spacer(1, 20))

    elements.append(
        Paragraph("Executive Summary", styles["Heading2"])
    )

    summary_table = [
        ["Compliance Score", summary.get("compliance_score", 0)],
        ["Risk Level", summary.get("risk_level", "Unknown")],
        ["Risk Score", summary.get("risk_score", 0)],
        ["Findings", summary.get("total_findings", 0)],
    ]

    table = Table(summary_table, colWidths=[180, 250])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )

    elements.append(table)
    elements.append(Spacer(1, 20))

    elements.append(
        Paragraph("Compliance Findings", styles["Heading2"])
    )

    if not findings:
        elements.append(
            Paragraph(
                "No material findings identified.",
                styles["Normal"],
            )
        )
    else:
        for i, finding in enumerate(findings, start=1):
            text = f"""
            <b>Finding {i}</b><br/>
            Category: {finding.get('category', 'N/A')}<br/>
            Severity: {finding.get('severity', 'N/A')}<br/>
            Status: {finding.get('status', 'N/A')}<br/>
            Issue: {finding.get('issue', 'N/A')}<br/>
            Evidence: {finding.get('evidence', 'N/A')}
            """
            elements.append(
                Paragraph(text, styles["Normal"])
            )
            elements.append(Spacer(1, 10))

    elements.append(
        Paragraph("Recommendations", styles["Heading2"])
    )

    if not recommendations:
        elements.append(
            Paragraph(
                "No recommendations returned.",
                styles["Normal"],
            )
        )
    else:
        for i, rec in enumerate(
            recommendations,
            start=1,
        ):
            text = f"""
            <b>Recommendation {i}</b><br/>
            Priority: {rec.get('priority', 'N/A')}<br/>
            Issue: {rec.get('issue', 'N/A')}<br/>
            Recommendation:
            {rec.get('recommendation', 'N/A')}
            """
            elements.append(
                Paragraph(text, styles["Normal"])
            )
            elements.append(Spacer(1, 10))

    doc.build(elements)
    buffer.seek(0)

    return buffer.getvalue()


def generate_docx_report(report: dict):
    document = Document()

    summary = report.get("executive_summary", {})
    findings = report.get("findings", [])
    recommendations = report.get("recommendations", [])

    document.add_heading(
        "AI Compliance & Document Review Report",
        level=1,
    )

    document.add_heading(
        "Document Information",
        level=2,
    )

    document.add_paragraph(
        f"File: {report.get('file_name', 'N/A')}"
    )
    document.add_paragraph(
        f"Framework: {report.get('framework', 'N/A')}"
    )
    document.add_paragraph(
        f"Document Type: {report.get('document_type', 'N/A')}"
    )
    document.add_paragraph(
        f"Generated: {report.get('created_at', 'N/A')}"
    )

    document.add_heading(
        "Executive Summary",
        level=2,
    )

    document.add_paragraph(
        f"Compliance Score: "
        f"{summary.get('compliance_score', 0)}%"
    )
    document.add_paragraph(
        f"Risk Level: "
        f"{summary.get('risk_level', 'Unknown')}"
    )
    document.add_paragraph(
        f"Risk Score: "
        f"{summary.get('risk_score', 0)}"
    )
    document.add_paragraph(
        f"Total Findings: "
        f"{summary.get('total_findings', 0)}"
    )

    document.add_heading(
        "Findings",
        level=2,
    )

    if not findings:
        document.add_paragraph(
            "No material findings identified."
        )
    else:
        for i, finding in enumerate(
            findings,
            start=1,
        ):
            p = document.add_paragraph()
            p.add_run(
                f"Finding {i}"
            ).bold = True

            p.add_run(
                f"\nCategory: "
                f"{finding.get('category', 'N/A')}"
            )
            p.add_run(
                f"\nSeverity: "
                f"{finding.get('severity', 'N/A')}"
            )
            p.add_run(
                f"\nStatus: "
                f"{finding.get('status', 'N/A')}"
            )
            p.add_run(
                f"\nIssue: "
                f"{finding.get('issue', 'N/A')}"
            )
            p.add_run(
                f"\nEvidence: "
                f"{finding.get('evidence', 'N/A')}"
            )

    document.add_heading(
        "Recommendations",
        level=2,
    )

    if not recommendations:
        document.add_paragraph(
            "No recommendations returned."
        )
    else:
        for i, rec in enumerate(
            recommendations,
            start=1,
        ):
            p = document.add_paragraph()
            p.add_run(
                f"Recommendation {i}"
            ).bold = True

            p.add_run(
                f"\nPriority: "
                f"{rec.get('priority', 'N/A')}"
            )
            p.add_run(
                f"\nIssue: "
                f"{rec.get('issue', 'N/A')}"
            )
            p.add_run(
                f"\nRecommendation: "
                f"{rec.get('recommendation', 'N/A')}"
            )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def generate_excel_report(report: dict):
    findings = report.get("findings", [])
    recommendations = report.get(
        "recommendations",
        [],
    )

    buffer = BytesIO()

    with pd.ExcelWriter(
        buffer,
        engine="openpyxl",
    ) as writer:

        summary = pd.DataFrame(
            [
                {
                    "File": report.get(
                        "file_name"
                    ),
                    "Framework": report.get(
                        "framework"
                    ),
                    "Document Type": report.get(
                        "document_type"
                    ),
                    "Risk Level": report[
                        "executive_summary"
                    ].get("risk_level"),
                    "Compliance Score": report[
                        "executive_summary"
                    ].get(
                        "compliance_score"
                    ),
                    "Risk Score": report[
                        "executive_summary"
                    ].get(
                        "risk_score"
                    ),
                }
            ]
        )

        summary.to_excel(
            writer,
            sheet_name="Summary",
            index=False,
        )

        pd.DataFrame(findings).to_excel(
            writer,
            sheet_name="Findings",
            index=False,
        )

        pd.DataFrame(
            recommendations
        ).to_excel(
            writer,
            sheet_name="Recommendations",
            index=False,
        )

    buffer.seek(0)
    return buffer.getvalue()